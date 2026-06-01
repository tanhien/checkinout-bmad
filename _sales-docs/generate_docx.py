"""
Generate polished DOCX sales documents for Hotel Management System.
Removes private email, adds system flow diagrams and demo screenshots.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
SCR_DIR  = os.path.join(DOCS_DIR, "screenshots")

# ─── Color palette ───────────────────────────────────────────────────────────
BLUE       = RGBColor(0x1d, 0x4e, 0xd8)   # blue-700
BLUE_LIGHT = RGBColor(0xdb, 0xea, 0xff)   # blue-50
GRAY       = RGBColor(0x6b, 0x72, 0x80)   # gray-500
DARK       = RGBColor(0x11, 0x18, 0x27)   # gray-900
WHITE      = RGBColor(0xff, 0xff, 0xff)
ORANGE     = RGBColor(0xea, 0x58, 0x0c)   # orange-600
GREEN      = RGBColor(0x16, 0xa3, 0x4a)   # green-600
RED        = RGBColor(0xdc, 0x26, 0x26)   # red-600

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def new_doc(title: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    return doc


def heading(doc, text: str, level: int = 1, color=BLUE):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Arial"
    run.font.bold = True
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc, text: str, bold=False, color=None, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = "Arial"
    run.font.size  = Pt(10)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    return p


def bullet(doc, text: str, level=0):
    style_name = "List Bullet" if level == 0 else "List Bullet 2"
    try:
        p = doc.add_paragraph(style=style_name)
    except Exception:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)


def info_box(doc, text: str, bg="DBEAFE", border_color="1D4ED8"):
    """Shaded callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()


def add_screenshot(doc, filename: str, caption: str, width_cm=14):
    path = os.path.join(SCR_DIR, filename)
    if not os.path.exists(path):
        body(doc, f"[Screenshot: {caption}]", color=GRAY)
        return
    doc.add_picture(path, width=Cm(width_cm))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = GRAY
    cap.paragraph_format.space_after = Pt(10)


def simple_table(doc, headers: list, rows: list, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, "1D4ED8")
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.name  = "Arial"
        run.font.size  = Pt(9)
    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = "F9FAFB" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            c.paragraphs[0].clear()
            run = c.paragraphs[0].add_run(str(cell_text))
            run.font.name = "Arial"
            run.font.size = Pt(9)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return tbl


def flow_box(doc, steps: list, title=""):
    """Render a horizontal flow as a table with arrows."""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(6)

    cols = len(steps) * 2 - 1
    tbl  = doc.add_table(rows=1, cols=cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Alternate: step cell, arrow cell
    for i, step in enumerate(steps):
        ci = i * 2
        cell = tbl.rows[0].cells[ci]
        set_cell_bg(cell, "DBEAFE")
        cell.width = Cm(2.8)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(step)
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = BLUE

        if ci + 1 < cols:
            arrow = tbl.rows[0].cells[ci + 1]
            arrow.width = Cm(0.5)
            set_cell_bg(arrow, "FFFFFF")
            pa = arrow.paragraphs[0]
            pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pa.add_run("→").font.size = Pt(11)

    # Remove all borders except between flow boxes
    doc.add_paragraph()


def contact_placeholder(doc):
    info_box(
        doc,
        "📧  Liên hệ tư vấn & demo:  [email@khachsan.vn]  ·  [SĐT liên hệ]\n"
        "🌐  Website:  [website.khachsan.vn]",
        bg="FEF9C3",
        border_color="CA8A04",
    )


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 — Tổng quan sản phẩm
# ═══════════════════════════════════════════════════════════════════════════════

def doc1_overview():
    doc = new_doc("Tổng Quan Sản Phẩm")

    # Cover header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HOTEL MANAGEMENT SYSTEM")
    run.font.name  = "Arial"
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Hệ Thống Quản Lý Khách Sạn Toàn Diện · Phiên bản 1.0 · 2026")
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    r2.font.color.rgb = GRAY

    info_box(doc,
        "Dành cho: Chủ khách sạn · Giám đốc vận hành · Nhà đầu tư\n"
        "Một nền tảng — Ba giao diện — Toàn bộ vận hành khách sạn trong tầm tay.",
        bg="DBEAFE")

    add_screenshot(doc, "booking-portal-home.png", "Hình 1: Trang đặt phòng trực tuyến — Lạc Hồng Boutique Hotel & Spa")

    # ── Vấn đề ──
    heading(doc, "1. Vấn Đề Thị Trường", 1)
    simple_table(doc,
        ["Tình trạng hiện tại", "Hậu quả cho khách sạn"],
        [
            ["Phần mềm PMS ngoại (Opera, Cloudbeds) $300–$800/tháng", "Chi phí không phù hợp với khách sạn vừa và nhỏ"],
            ["Quản lý bằng Excel, Zalo, sổ tay", "Sai sót, mất dữ liệu, không có báo cáo tổng hợp"],
            ["Không có kiosk tự phục vụ", "Hàng đợi giờ cao điểm, cần nhiều nhân sự lễ tân"],
            ["Không có portal đặt phòng riêng", "Phụ thuộc 100% OTA (Booking.com/Agoda), mất phí 15–25%"],
        ],
        col_widths=[9, 7]
    )

    # ── Ba hệ thống ──
    heading(doc, "2. Giải Pháp HMS — Ba Hệ Thống Liên Thông", 1)

    flow_box(doc,
        ["Booking Portal\n(Khách đặt online)", "Kiosk App\n(Check-in/out kiosk)", "Staff Web\n(Vận hành nội bộ)"],
        title="Kiến trúc tổng quan — 3 giao diện kết nối cùng 1 cơ sở dữ liệu:"
    )
    body(doc, "")

    # Booking portal
    heading(doc, "2.1  Booking Portal — Đặt Phòng Trực Tuyến", 2)
    body(doc, "Trang web đặt phòng chuyên nghiệp tích hợp vào website khách sạn. Khách tìm phòng, đặt phòng, thanh toán — không qua bên thứ ba.")
    for b in ["Giao diện song ngữ Việt / Anh (tự động theo trình duyệt)",
              "Tìm phòng theo ngày, hiển thị giá và tình trạng còn phòng real-time",
              "Quy trình đặt phòng 4 bước trực quan, hỗ trợ mã khuyến mãi (promo code)",
              "Email xác nhận tự động kèm QR code để check-in tại kiosk",
              "Trang quản lý đặt phòng cho khách (xem, hủy, tra cứu không cần đăng nhập)"]:
        bullet(doc, b)

    add_screenshot(doc, "booking-portal-rooms.png", "Hình 2: Danh sách phòng với giá real-time và bộ lọc")

    # Kiosk
    heading(doc, "2.2  Kiosk App — Tự Phục Vụ 24/7", 2)
    body(doc, "Ứng dụng fullscreen chạy trên Android tablet hoặc Windows PC tại quầy lễ tân. Khách tự check-in/out không cần nhân viên.")
    for b in ["Check-in bằng mã xác nhận hoặc quét QR code",
              "Tự động xếp phòng và hiển thị số phòng",
              "Đặt phòng trực tiếp (walk-in) ngay tại kiosk với thanh toán tích hợp",
              "Check-out, xem hóa đơn và thanh toán tại kiosk",
              "Nút gọi nhân viên khẩn cấp — cảnh báo real-time trên màn hình lễ tân"]:
        bullet(doc, b)

    add_screenshot(doc, "kiosk-home.png", "Hình 3: Màn hình chờ kiosk — 3 chức năng chính", width_cm=8)

    # Staff web
    heading(doc, "2.3  Staff Management Web — Trung Tâm Vận Hành", 2)
    body(doc, "Ứng dụng web dành cho toàn bộ nhân sự khách sạn. Mỗi vai trò chỉ thấy đúng những gì cần thiết cho công việc của mình.")
    for b in ["Dashboard real-time: arrivals, departures, công suất, cảnh báo kiosk",
              "Quản lý booking từ mọi kênh: trực tiếp, kiosk, online, OTA",
              "Housekeeping Kanban: phân công và theo dõi trạng thái dọn phòng",
              "Folio & Thanh toán: phụ phí, discount, ghi nhận thanh toán, xuất PDF",
              "Báo cáo: Occupancy, Doanh thu (RevPAR, ADR), Kênh đặt phòng, Arrivals/Departures"]:
        bullet(doc, b)

    add_screenshot(doc, "staff-dashboard.png", "Hình 4: Dashboard lễ tân — tổng quan real-time toàn khách sạn")

    # ── Flow hệ thống ──
    heading(doc, "3. Luồng Hệ Thống Tổng Quan", 1)

    heading(doc, "3.1  Flow Đặt Phòng Online → Check-in Kiosk", 2)
    flow_box(doc, ["Khách vào\nBooking Portal", "Tìm phòng\n& đặt", "Xác nhận\n& email QR", "Đến khách sạn\nquét QR tại kiosk", "Nhận số phòng\n< 90 giây"])
    body(doc, "Luồng này hoàn toàn tự động — không cần nhân viên can thiệp.")

    heading(doc, "3.2  Flow Staff Check-in Thủ Công", 2)
    flow_box(doc, ["Tìm booking\nCONFIRMED", "Xác nhận\nthông tin khách", "Gán phòng\nkhả dụng", "Nhấn\nCheck-in", "Booking →\nCHECKED_IN"])
    body(doc, "Dùng khi khách cần hỗ trợ hoặc không có kiosk.")

    heading(doc, "3.3  Flow Check-out & Thanh Toán", 2)
    flow_box(doc, ["Mở Folio\ncủa booking", "Kiểm tra\nphụ phí", "Ghi nhận\nthanh toán", "Nhấn\nCheck-out", "Settle Folio\n& xuất PDF"])
    body(doc, "Phòng tự động chuyển sang DIRTY sau check-out — Housekeeping nhận thông báo.")

    heading(doc, "3.4  Flow Housekeeping", 2)
    flow_box(doc, ["Check-out\nxong", "Phòng →\nDIRTY", "Manager\nphân công", "HK dọn →\nCLEANING", "Hoàn thành →\nCLEAN"])

    add_screenshot(doc, "staff-housekeeping.png", "Hình 5: Housekeeping Kanban board — quản lý dọn phòng real-time")

    # ── Lợi thế ──
    heading(doc, "4. Lợi Thế Cạnh Tranh", 1)
    simple_table(doc,
        ["", "HMS", "Opera / Cloudbeds", "Phần mềm VN nội địa"],
        [
            ["Chi phí license/tháng", "Không có", "$300–$800", "$50–$200"],
            ["Kiosk tự phục vụ", "✅ Có sẵn", "Addon riêng", "Không"],
            ["Portal đặt phòng riêng", "✅ Có sẵn", "Tích hợp cơ bản", "Không"],
            ["Real-time toàn hệ thống", "✅ Socket.io", "Có", "Hạn chế"],
            ["Tùy biến theo yêu cầu", "✅ Toàn quyền", "Không", "Hạn chế"],
            ["Tiếng Việt đầy đủ", "✅ 100%", "Hạn chế", "Có"],
            ["Thời gian triển khai", "✅ 2–4 tuần", "2–6 tháng", "1–3 tháng"],
        ],
        col_widths=[5, 3.5, 3.5, 4]
    )

    # ── Phù hợp ──
    heading(doc, "5. Phù Hợp Với Loại Hình Nào?", 1)
    for b in ["Khách sạn 10–200 phòng", "Boutique hotel, resort, serviced apartment, hostel",
              "Chuỗi khách sạn nhỏ (multi-property)", "Khách sạn đang chuyển đổi số từ Excel/thủ công"]:
        bullet(doc, b)

    add_divider(doc)
    contact_placeholder(doc)

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2 — Danh sách tính năng
# ═══════════════════════════════════════════════════════════════════════════════

def doc2_features():
    doc = new_doc("Danh Sách Tính Năng")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DANH SÁCH TÍNH NĂNG ĐẦY ĐỦ")
    run.font.name = "Arial"; run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = BLUE

    info_box(doc, "Hotel Management System v1.0 — 2026\nBao gồm: Staff Web · Booking Portal · Kiosk App · API Backend", bg="DBEAFE")

    sections = [
        ("📟  KIOSK APP", [
            ("Check-in tự phục vụ", [
                "Nhập mã xác nhận 6 ký tự hoặc quét QR code từ email",
                "Xác minh tên khách (chuẩn hóa dấu tiếng Việt)",
                "Tự động xếp phòng tối ưu (optimistic locking, thử lại 5 lần)",
                "Màn hình thành công hiển thị số phòng",
            ]),
            ("Đặt phòng walk-in tại kiosk", [
                "Chọn ngày check-in/out và số người",
                "Xem danh sách loại phòng với giá real-time",
                "Nhập thông tin khách",
                "Thanh toán demo (2 giây) → tự động check-in và xếp phòng",
            ]),
            ("Check-out tại kiosk", [
                "Tra cứu theo số phòng hoặc mã xác nhận",
                "Xem folio và tổng số tiền",
                "Thanh toán và nhận xác nhận check-out",
            ]),
            ("Tiện ích khác", [
                "Nút 'Gọi nhân viên' — cảnh báo realtime qua WebSocket",
                "Màn hình chờ idle tự động sau 5 phút",
                "Chuyển ngôn ngữ VI / EN tại header",
                "Chạy fullscreen PWA trên Android tablet & Windows PC",
            ]),
        ]),
        ("🖥️  STAFF MANAGEMENT WEB", [
            ("Dashboard lễ tân", [
                "Arrivals / Departures / Đang ở hôm nay",
                "Cảnh báo kiosk real-time (âm thanh beep + nút dismiss)",
                "Widget tình trạng phòng (CLEAN/DIRTY/OCCUPIED/MAINTENANCE)",
                "Biểu đồ dự báo công suất 7–14 ngày tới",
                "Auto-refresh 60 giây",
            ]),
            ("Dashboard buồng phòng (Housekeeping role)", [
                "Chỉ thấy phòng được phân công cho nhân viên đó",
                "Trạng thái và loại phòng từng thẻ",
            ]),
            ("Quản lý Booking", [
                "Danh sách booking với tìm kiếm full-text và bộ lọc đa chiều",
                "Lọc theo: trạng thái, kênh, loại phòng, ngày check-in/out",
                "Tạo booking mới: tìm khách có sẵn hoặc tạo mới, chọn rate plan",
                "Chi tiết booking: thông tin đầy đủ, lịch sử thay đổi (audit log)",
                "Gán phòng / bỏ gán phòng",
                "Check-in thủ công (sau giờ check-in cấu hình)",
                "Check-out thủ công",
                "Hủy booking (ghi lý do, soft delete)",
                "Đánh dấu No-show",
            ]),
            ("Quản lý Phòng", [
                "Grid phòng theo loại, lọc theo trạng thái",
                "Đổi trạng thái: CLEAN / DIRTY / MAINTENANCE",
                "6 trạng thái phòng: CLEAN · DIRTY · CLEANING · INSPECTED · OCCUPIED · MAINTENANCE",
            ]),
            ("Housekeeping Kanban", [
                "Board kéo-thả: DIRTY → CLEANING → CLEAN → INSPECTED",
                "Phân công phòng cho từng nhân viên buồng phòng",
                "Tracking thời gian bắt đầu dọn",
                "Real-time qua WebSocket",
            ]),
            ("Khách hàng (Guest profiles)", [
                "Tìm kiếm full-text: tên, email, SĐT, quốc tịch",
                "Lịch sử toàn bộ đặt phòng",
                "CMND/Hộ chiếu mã hóa AES-256-GCM",
                "Chương trình loyalty (điểm tích lũy)",
            ]),
            ("Folio & Thanh toán", [
                "Tổng quan folio: stat cards + bảng chia theo nhóm",
                "Cảnh báo folio quá hạn (đã check-out chưa settle)",
                "Thêm phụ phí dịch vụ (từ danh mục dịch vụ cấu hình)",
                "Thêm discount: % hoặc số tiền cố định (>10% cần Manager)",
                "Void phụ phí (ghi lý do, không xóa vật lý)",
                "Ghi nhận thanh toán: Tiền mặt / Thẻ / Chuyển khoản",
                "Settle folio sau khi thanh toán đủ",
                "Reopen folio đã settle (Manager only)",
                "Xuất hóa đơn PDF tiếng Việt (font Arial, MST, tài khoản ngân hàng)",
            ]),
            ("Báo cáo", [
                "Occupancy: % công suất theo ngày/tuần/tháng, so sánh kỳ trước",
                "Revenue: tổng doanh thu, RevPAR, ADR, biểu đồ cột",
                "Arrivals & Departures: danh sách theo ngày, lọc theo khoảng thời gian",
                "Channels: phân tích theo kênh đặt phòng (Direct/Kiosk/Online/OTA...)",
            ]),
            ("Cấu hình Admin", [
                "Thông tin khách sạn: tên, địa chỉ, MST, tài khoản ngân hàng",
                "Loại phòng: CRUD + upload ảnh (tối đa 10 ảnh, WebP, Cloudflare R2)",
                "Phòng vật lý: số phòng, tầng, gán loại phòng",
                "Gói giá (Rate Plans): giá ngày thường/cuối tuần, % discount, hoàn tiền",
                "Dịch vụ & phụ phí: danh mục dịch vụ bán thêm",
                "Thuế (VAT/Service charge): tự động áp vào folio",
                "Mã khuyến mãi: % hoặc số tiền cố định, giới hạn sử dụng, hết hạn",
                "Nhân viên: tạo tài khoản, onboarding email, reset password",
                "Email template: Handlebars HTML, preview trước khi lưu",
            ]),
        ]),
        ("🌐  BOOKING PORTAL", [
            ("Trang công khai", [
                "Homepage: hero ảnh, search widget, flash deals, phòng nổi bật, gallery",
                "Trang Giới thiệu: câu chuyện khách sạn, giải thưởng, chính sách, vị trí",
                "Trang Tiện nghi: hồ bơi, spa, nhà hàng, gym, dịch vụ",
                "Danh sách phòng: sort, slideshow ảnh, badge tiện nghi",
                "Chi tiết phòng: lightbox ảnh, ISR, widget đặt phòng tích hợp",
                "Privacy Policy & Terms of Use",
            ]),
            ("Funnel đặt phòng 4 bước", [
                "Bước 1: Xác nhận phòng, nhập mã promo code (validate real-time)",
                "Bước 2: Thông tin khách (tự điền từ tài khoản nếu đã đăng nhập)",
                "Bước 3: Xem lại đầy đủ trước khi xác nhận",
                "Bước 4: Xác nhận, kiểm tra lại tình trạng phòng, email tự động",
            ]),
            ("Trang xác nhận booking", [
                "QR code server-side (render thành base64 PNG)",
                "Thêm vào Google Calendar",
                "Tải file ICS (iCal)",
                "In trang xác nhận",
                "Link quản lý đặt phòng",
            ]),
            ("Guest auth & My Bookings", [
                "Đăng ký / Đăng nhập / Quên mật khẩu / Reset password",
                "JWT httpOnly cookie (7 ngày / 30 ngày)",
                "My Bookings: 3 tab Sắp tới / Đã ở / Đã hủy",
                "Hủy booking trong cửa sổ hủy miễn phí",
                "Lookup booking không cần đăng nhập (mã xác nhận + email)",
            ]),
            ("Kỹ thuật & SEO", [
                "next-intl: URL prefix /vi/... và /en/... tự động",
                "generateMetadata, Schema.org, sitemap.xml, robots.txt",
                "ISR (Incremental Static Regeneration) cho trang phòng",
                "Email qua Resend (silent fail nếu chưa cấu hình)",
            ]),
        ]),
    ]

    for sec_title, subsections in sections:
        heading(doc, sec_title, 1)
        for sub_title, items in subsections:
            heading(doc, sub_title, 2, color=DARK)
            for item in items:
                bullet(doc, item)
        doc.add_paragraph()

    # Screenshots section
    heading(doc, "📸  Demo Screenshots", 1)
    add_screenshot(doc, "staff-bookings.png", "Hình 6: Danh sách booking — lọc đa chiều, tạo mới")
    add_screenshot(doc, "staff-folio.png", "Hình 7: Tổng quan Folio & Thanh toán — cảnh báo quá hạn")
    add_screenshot(doc, "booking-portal-book.png", "Hình 8: Funnel đặt phòng 4 bước — bước xác nhận")

    add_divider(doc)
    contact_placeholder(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 3 — Kịch bản demo
# ═══════════════════════════════════════════════════════════════════════════════

def doc3_demo():
    doc = new_doc("Kịch Bản Demo")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KỊCH BẢN DEMO")
    run.font.name = "Arial"; run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Dành cho đội ngũ Sales · Thời lượng: 30–45 phút")
    r2.font.name = "Arial"; r2.font.size = Pt(11); r2.font.color.rgb = GRAY

    info_box(doc, "Mục tiêu: Khách hàng thấy toàn bộ luồng từ đặt phòng online → check-in kiosk → vận hành staff → thanh toán.", bg="DCFCE7")

    # Checklist
    heading(doc, "Chuẩn Bị Trước Demo", 1)
    for b in [
        "Chạy: docker compose up -d (PostgreSQL + Redis)",
        "Chạy: pnpm dev (3 app khởi động song song)",
        "Mở 3 tab: localhost:3001 (Staff) · localhost:3000/vi (Portal) · localhost:3002 (Kiosk)",
        "Đăng nhập Staff Web: admin@demo.hotel / Admin123!",
        "Kiểm tra: có ít nhất 1 phòng trạng thái CLEAN",
        "Tắt thông báo điện thoại / chế độ 'Không làm phiền'",
    ]:
        bullet(doc, b)

    # Flow tổng quát
    heading(doc, "Luồng Demo 30 Phút", 1)
    flow_box(doc, ["Booking Portal\n(5 phút)", "Staff Dashboard\n(5 phút)", "Kiosk Check-in\n(5 phút)", "Housekeeping\n(5 phút)", "Folio & Báo cáo\n(5 phút)", "Q&A / Chốt\n(10 phút)"])

    add_screenshot(doc, "booking-portal-home.png", "Màn hình 1: Booking Portal — điểm khởi đầu demo")

    # Phần 1
    heading(doc, "Phần 1 — Booking Portal (5 phút)", 1)
    info_box(doc, '💬 Mở đầu: "Hãy bắt đầu từ góc nhìn của khách hàng — đây là trang đặt phòng riêng của khách sạn, không qua Booking.com, không mất phí hoa hồng 15–25%."', bg="FEF9C3")
    heading(doc, "Các bước thao tác:", 2)
    for step in [
        "Mở localhost:3000/vi — chỉ vào hero banner và tên khách sạn",
        "Nhập ngày check-in/out vào search widget → Click 'Tìm phòng trống'",
        "Chọn 1 phòng → Xem trang chi tiết (ảnh slideshow, tiện nghi, giá)",
        "Click 'Đặt ngay' → Đi qua 4 bước đặt phòng",
        "Bước cuối: Chỉ QR code — 'QR này dùng để check-in tại kiosk, sẽ demo ngay sau'",
    ]:
        bullet(doc, step)
    info_box(doc, '✨ Điểm nhấn: "Khách đặt trực tiếp — anh/chị nhận 100% doanh thu."', bg="DCFCE7")

    # Phần 2
    heading(doc, "Phần 2 — Staff Web: Dashboard (5 phút)", 1)
    info_box(doc, '💬 Chuyển tiếp: "Ngay khi khách vừa đặt, nhân viên thấy ngay — không cần F5, không cần báo cáo buổi sáng."', bg="FEF9C3")
    for step in [
        "Chỉ vào Dashboard: số khách đến/đi hôm nay, tình trạng phòng",
        "Mở tab Booking → tìm booking vừa tạo → xem chi tiết",
        "Chỉ Audit Log — 'Mọi thay đổi đều được ghi lại: ai, lúc nào, làm gì'",
        "Click 'Check-in thủ công' — 'Nhân viên vẫn hỗ trợ khi khách cần'",
    ]:
        bullet(doc, step)

    add_screenshot(doc, "staff-dashboard.png", "Màn hình 2: Staff Dashboard — tổng quan real-time")
    add_screenshot(doc, "staff-bookings.png", "Màn hình 3: Danh sách booking — tất cả kênh đặt phòng")

    # Phần 3 — Kiosk
    heading(doc, "Phần 3 — Kiosk: Check-in Tự Phục Vụ (5 phút)", 1)
    info_box(doc, '💬 Chuyển tiếp: "Bây giờ hãy đóng vai khách hàng — đến kiosk, quét QR, nhận phòng trong dưới 90 giây."', bg="FEF9C3")
    flow_box(doc, ["Quét QR\nhoặc nhập mã", "Xác nhận\ntên khách", "Tự động\nxếp phòng", "Màn hình\nthành công"])
    for step in [
        "Mở localhost:3002 — Chỉ giao diện fullscreen 3 nút",
        "Click 'Tôi đã đặt phòng' → Nhập mã xác nhận từ email",
        "Xác nhận tên → Hệ thống tự xếp phòng",
        "Quay lại Staff Dashboard — chỉ thông báo check-in xuất hiện real-time",
    ]:
        bullet(doc, step)

    add_screenshot(doc, "kiosk-home.png", "Màn hình 4: Kiosk — giao diện tự phục vụ", width_cm=9)
    info_box(doc, '✨ Điểm nhấn: "Không cần nhân viên — khách tự check-in lúc 2 giờ sáng cũng không vấn đề."', bg="DCFCE7")

    # Phần 4 — Housekeeping
    heading(doc, "Phần 4 — Housekeeping & Quản Lý Phòng (5 phút)", 1)
    for step in [
        "Mở tab Housekeeping → Chỉ Kanban board",
        "Kéo thẻ phòng từ DIRTY sang CLEANING → sang CLEAN",
        "Mở tab Phòng → Grid màu sắc tình trạng từng phòng",
        "Chỉ: 'Màu xanh = sạch sẵn sàng, xanh đậm = có khách, đỏ = bảo trì'",
    ]:
        bullet(doc, step)
    add_screenshot(doc, "staff-housekeeping.png", "Màn hình 5: Housekeeping Kanban board")
    add_screenshot(doc, "staff-rooms.png", "Màn hình 6: Grid trạng thái phòng")

    # Phần 5 — Folio & Báo cáo
    heading(doc, "Phần 5 — Folio & Báo Cáo (5 phút)", 1)
    flow_box(doc, ["Check-out\nxong", "Mở Folio", "Thêm phụ phí\n(nếu có)", "Ghi nhận\nthanh toán", "Settle\n& xuất PDF"])
    for step in [
        "Mở Folio & Thanh toán — chỉ dashboard tổng quan, stat cards",
        "Click vào 1 folio → Xem phụ phí, thêm discount",
        "Ghi nhận thanh toán → Settle",
        "Download PDF hóa đơn → Chỉ hóa đơn tiếng Việt đầy đủ",
        "Mở Báo cáo Doanh thu → Chỉ biểu đồ, RevPAR, ADR",
    ]:
        bullet(doc, step)
    add_screenshot(doc, "staff-folio.png", "Màn hình 7: Folio & Thanh toán")
    add_screenshot(doc, "staff-reports-occupancy.png", "Màn hình 8: Báo cáo công suất")

    add_divider(doc)
    contact_placeholder(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 4 — Bảng giá
# ═══════════════════════════════════════════════════════════════════════════════

def doc4_pricing():
    doc = new_doc("Bảng Giá")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BẢNG GIÁ TRIỂN KHAI")
    run.font.name = "Arial"; run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Hotel Management System · Cập nhật 2026 · Đơn vị: VNĐ · Chưa bao gồm VAT")
    r2.font.name = "Arial"; r2.font.size = Pt(10); r2.font.color.rgb = GRAY

    info_box(doc, "Giá có thể điều chỉnh theo quy mô và yêu cầu cụ thể. Liên hệ để nhận báo giá chính xác nhất.", bg="FEF9C3")

    heading(doc, "1. Gói Dịch Vụ", 1)
    simple_table(doc,
        ["Hạng mục", "Starter\n(≤ 30 phòng)", "Professional\n(31–100 phòng)", "Enterprise\n(> 100 phòng)"],
        [
            ["Phí triển khai một lần", "35.000.000 đ", "65.000.000 đ", "Báo giá riêng"],
            ["Phí hỗ trợ hàng tháng", "2.000.000 đ", "3.500.000 đ", "Theo hợp đồng"],
            ["Staff Management Web", "✅", "✅", "✅"],
            ["Booking Portal (VI/EN)", "✅", "✅ + domain riêng", "✅"],
            ["Kiosk App", "1 thiết bị", "2 thiết bị", "Không giới hạn"],
            ["Payment gateway", "1 (VNPay/MoMo)", "2 gateway", "Theo yêu cầu"],
            ["Custom email template", "—", "✅", "✅"],
            ["Tùy chỉnh giao diện brand", "—", "Cơ bản", "Nâng cao"],
            ["Multi-property", "—", "—", "✅"],
            ["Training", "4h online", "8h (trực tiếp+online)", "Onsite"],
            ["Bảo hành", "6 tháng", "12 tháng", "Theo hợp đồng"],
            ["SLA hỗ trợ", "Giờ hành chính", "8x5, phản hồi 4h", "24/7, phản hồi 1h"],
        ],
        col_widths=[5, 3.5, 3.5, 3.5]
    )

    heading(doc, "2. Dịch Vụ Bổ Sung (Add-on)", 1)
    simple_table(doc,
        ["Dịch vụ", "Giá"],
        [
            ["Thêm 1 thiết bị kiosk", "8.000.000 đ/thiết bị"],
            ["Tích hợp thêm 1 payment gateway", "5.000.000 đ"],
            ["Tùy chỉnh giao diện nâng cao", "10.000.000 – 20.000.000 đ"],
            ["Tích hợp channel manager (SiteMinder...)", "15.000.000 đ"],
            ["Tích hợp phần mềm kế toán (MISA, Fast)", "12.000.000 đ"],
            ["Migration dữ liệu từ hệ thống cũ", "5.000.000 – 15.000.000 đ"],
            ["Training bổ sung", "1.500.000 đ/giờ"],
            ["Phát triển tính năng theo yêu cầu", "Báo giá theo scope"],
        ],
        col_widths=[10, 6]
    )

    heading(doc, "3. So Sánh Chi Phí 3 Năm (Khách Sạn 50 Phòng)", 1)
    simple_table(doc,
        ["", "HMS Professional", "Cloudbeds Lodging"],
        [
            ["Phí setup / triển khai", "65.000.000 đ", "~30.000.000 đ (onboarding)"],
            ["Phí hàng tháng (36 tháng)", "3.500.000 × 36 = 126.000.000 đ", "~$280 × 36 = ~252.000.000 đ"],
            ["Tổng 3 năm", "191.000.000 đ", "~282.000.000 đ"],
            ["Tiết kiệm", "—", "~91.000.000 đ"],
        ],
        col_widths=[6, 6, 5]
    )
    body(doc, "* Chưa tính tiết kiệm hoa hồng OTA khi có booking portal riêng (~8–15 triệu đ/tháng).", color=GRAY)

    heading(doc, "4. Điều Kiện & Chính Sách", 1)
    for b in [
        "Thanh toán: 50% khi ký hợp đồng, 50% khi bàn giao",
        "Thời gian triển khai: 2–4 tuần kể từ ngày ký",
        "Bảo hành: Lỗi hệ thống được fix miễn phí trong thời gian bảo hành",
        "Dữ liệu: Khách hàng sở hữu toàn bộ dữ liệu, có thể xuất ra bất kỳ lúc nào",
        "Pilot: Có thể thảo luận gói pilot 1 tháng cho khách hàng chưa quyết định",
    ]:
        bullet(doc, b)

    heading(doc, "Để Nhận Báo Giá Chính Xác, Vui Lòng Cung Cấp:", 1)
    for b in ["Số phòng", "Loại hình (hotel/hostel/resort/serviced apartment)",
              "Hệ thống đang dùng (nếu có)", "Yêu cầu đặc biệt"]:
        bullet(doc, b)

    add_divider(doc)
    contact_placeholder(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 5 — ROI
# ═══════════════════════════════════════════════════════════════════════════════

def doc5_roi():
    doc = new_doc("ROI Business Case")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ROI & BUSINESS CASE")
    run.font.name = "Arial"; run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Phân Tích Lợi Nhuận Đầu Tư — Hotel Management System")
    r2.font.name = "Arial"; r2.font.size = Pt(11); r2.font.color.rgb = GRAY

    info_box(doc, "Ví dụ tính toán dựa trên: Khách sạn 50 phòng, ADR 800.000 đ, công suất 70%, 30 ngày/tháng.", bg="DBEAFE")

    heading(doc, "1. Tiết Kiệm Chi Phí Hoa Hồng OTA", 1)
    simple_table(doc,
        ["Kịch bản", "Trước HMS", "Sau HMS (6 tháng)"],
        [
            ["% booking qua OTA", "80%", "50%"],
            ["% booking qua portal riêng", "0%", "30%"],
            ["Hoa hồng OTA trung bình", "20%", "20%"],
            ["Doanh thu/tháng (50 phòng, 70% occ, 800K ADR)", "840.000.000 đ", "840.000.000 đ"],
            ["Hoa hồng OTA phải trả/tháng", "134.400.000 đ", "84.000.000 đ"],
            ["Tiết kiệm hoa hồng/tháng", "—", "50.400.000 đ"],
            ["Tiết kiệm hoa hồng/năm", "—", "604.800.000 đ"],
        ],
        col_widths=[7.5, 4.5, 4.5]
    )

    heading(doc, "2. Tiết Kiệm Chi Phí Nhân Sự (Kiosk tự phục vụ)", 1)
    simple_table(doc,
        ["Kịch bản", "Trước HMS", "Sau HMS"],
        [
            ["Nhân viên lễ tân ca đêm", "2 người", "1 người"],
            ["Lương/người/tháng", "8.000.000 đ", "8.000.000 đ"],
            ["Tiết kiệm nhân sự/tháng", "—", "8.000.000 đ"],
            ["Tiết kiệm nhân sự/năm", "—", "96.000.000 đ"],
        ],
        col_widths=[8, 4.5, 4.5]
    )

    heading(doc, "3. Tổng Hợp ROI", 1)
    simple_table(doc,
        ["Hạng mục", "Năm 1", "Năm 2", "Năm 3"],
        [
            ["Tiết kiệm hoa hồng OTA", "604.800.000 đ", "604.800.000 đ", "604.800.000 đ"],
            ["Tiết kiệm nhân sự", "96.000.000 đ", "96.000.000 đ", "96.000.000 đ"],
            ["Chi phí HMS (Professional)", "-107.000.000 đ", "-42.000.000 đ", "-42.000.000 đ"],
            ["Lợi nhuận ròng", "593.800.000 đ", "658.800.000 đ", "658.800.000 đ"],
            ["ROI", "555%", "1.570%", "2.585%"],
        ],
        col_widths=[6.5, 3.5, 3.5, 3.5]
    )
    body(doc, "Năm 1: 65M (triển khai) + 42M (12 tháng hỗ trợ) = 107M. Năm 2+: chỉ 42M/năm.", color=GRAY)

    heading(doc, "4. Thời Gian Hoàn Vốn", 1)
    flow_box(doc, ["Ký hợp đồng\nTháng 1", "Triển khai\nTháng 2–4", "Bắt đầu tiết kiệm\nTháng 5", "Hoàn vốn\nTháng 7–8", "Lợi nhuận ròng\nTháng 9+"])
    body(doc, "Thời gian hoàn vốn trung bình: 3–5 tháng sau khi đi live.")

    heading(doc, "5. Lợi Ích Không Định Lượng Được", 1)
    for b in [
        "Nâng cao trải nghiệm khách hàng — ít chờ đợi, check-in nhanh",
        "Báo cáo real-time giúp quyết định nhanh, chính xác hơn",
        "Giảm sai sót do thủ công (overbooking, nhập liệu sai, mất thông tin)",
        "Dữ liệu khách hàng tập trung — cơ sở cho chương trình loyalty",
        "Hình ảnh chuyên nghiệp khi có kiosk hiện đại tại khách sạn",
    ]:
        bullet(doc, b)

    add_divider(doc)
    contact_placeholder(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 6 — FAQ
# ═══════════════════════════════════════════════════════════════════════════════

def doc6_faq():
    doc = new_doc("FAQ Bán Hàng")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CÂU HỎI THƯỜNG GẶP (FAQ)")
    run.font.name = "Arial"; run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = BLUE

    faqs = [
        ("🏗️  Triển khai & Kỹ thuật", [
            ("Triển khai mất bao lâu?",
             "2–4 tuần sau khi ký hợp đồng. Gồm: cài đặt hạ tầng (3–5 ngày), khai báo dữ liệu khách sạn (2–3 ngày), training nhân viên (1–2 ngày), kiểm tra và go-live."),
            ("Hệ thống chạy trên đâu?",
             "Cloud VPS (khuyến nghị), Railway/Render, hoặc server tại chỗ. Cần: 2 vCPU, 4GB RAM, 50GB SSD, PostgreSQL, Redis. Chi phí hosting ~200.000–500.000 đ/tháng."),
            ("Có cần internet liên tục không?",
             "Có. Kiosk và Staff web cần kết nối ổn định. Khuyến nghị: đường backup 4G/5G cho kiosk."),
            ("Hỗ trợ những payment gateway nào?",
             "Demo adapter có sẵn. Tích hợp thực tế: VNPay, MoMo, Stripe, ZaloPay — theo yêu cầu."),
            ("Có thể tích hợp với phần mềm kế toán không?",
             "Có — tích hợp MISA, Fast Accounting theo yêu cầu. Phí riêng."),
        ]),
        ("💰  Giá & Hợp đồng", [
            ("Có phí license hàng tháng không?",
             "Không có phí license. Chỉ có phí hỗ trợ kỹ thuật hàng tháng (2–3.5 triệu tùy gói). Tất cả tính năng đều có sẵn, không giới hạn user."),
            ("Nếu muốn thêm tính năng riêng thì sao?",
             "Báo giá theo scope. Hệ thống mã nguồn mở về phía khách hàng — có thể tự phát triển thêm nếu có đội IT."),
            ("Điều gì xảy ra nếu chúng tôi muốn dừng hợp đồng?",
             "Khách hàng sở hữu toàn bộ dữ liệu và mã nguồn. Có thể xuất dữ liệu bất kỳ lúc nào. Không bị khoá vendor."),
        ]),
        ("🏨  Vận hành", [
            ("Kiosk chạy trên thiết bị gì?",
             "Android tablet (Android 10+, màn hình 10\"+) hoặc Windows PC (Windows 10, màn hình cảm ứng). Khách hàng tự mua phần cứng, HMS tư vấn model phù hợp."),
            ("Nhân viên cần training bao lâu?",
             "Thường 4–8 giờ. Lễ tân/Manager thành thạo sau 1–2 ngày thực tế. Housekeeping chỉ cần 30 phút cho Kanban board."),
            ("Khi có sự cố, hỗ trợ như thế nào?",
             "Zalo Group 24/7 cho câu hỏi nhanh. Email/hotline theo SLA gói. Sự cố nghiêm trọng: phản hồi trong 1–4 giờ tùy gói."),
            ("Hệ thống xử lý giờ cao điểm (nhiều check-in cùng lúc) thế nào?",
             "Optimistic locking tránh double booking. Kiosk dùng retry 5 lần khi phòng bị lấy đồng thời. Hệ thống test thực tế với 50+ booking đồng thời."),
        ]),
        ("🔐  Bảo mật & Dữ liệu", [
            ("Dữ liệu khách có an toàn không?",
             "CMND/Hộ chiếu mã hóa AES-256-GCM. JWT httpOnly cookie. Mọi thay đổi ghi Audit Log. Không lưu thông tin thẻ tín dụng (xử lý qua payment gateway)."),
            ("Backup dữ liệu thế nào?",
             "Khuyến nghị: daily backup PostgreSQL (cron + Supabase/Railway auto backup). Ảnh phòng lưu trên Cloudflare R2 (99.99% durability)."),
        ]),
    ]

    for section_title, questions in faqs:
        heading(doc, section_title, 1)
        for q, a in questions:
            p = doc.add_paragraph()
            run = p.add_run(f"Q: {q}")
            run.font.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = BLUE
            p.paragraph_format.space_before = Pt(6)
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"A: {a}")
            run2.font.name = "Arial"
            run2.font.size = Pt(10)
            p2.paragraph_format.left_indent = Cm(0.5)
            p2.paragraph_format.space_after  = Pt(4)

    add_divider(doc)
    contact_placeholder(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 7 — Kế hoạch triển khai
# ═══════════════════════════════════════════════════════════════════════════════

def doc7_deployment():
    doc = new_doc("Kế Hoạch Triển Khai")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KẾ HOẠCH TRIỂN KHAI")
    run.font.name = "Arial"; run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Hotel Management System · Thời gian: 2–4 tuần từ ngày ký hợp đồng")
    r2.font.name = "Arial"; r2.font.size = Pt(11); r2.font.color.rgb = GRAY

    # Timeline flow
    heading(doc, "1. Lộ Trình Tổng Quan", 1)
    flow_box(doc, ["Ký hợp đồng\nNgày 1", "Hạ tầng\nNgày 1–5", "Khai báo dữ liệu\nNgày 6–10", "Training\nNgày 11–14", "UAT & Go-live\nNgày 14–21"])

    heading(doc, "2. Chi Tiết Từng Giai Đoạn", 1)
    phases = [
        ("Giai Đoạn 1 — Hạ Tầng (Ngày 1–5)", [
            "Cài đặt VPS / server theo yêu cầu",
            "Cài đặt PostgreSQL, Redis, NGINX, SSL certificate",
            "Deploy 3 ứng dụng: Staff Web, Booking Portal, Kiosk App",
            "Cấu hình biến môi trường: database, JWT secrets, email API, R2 storage",
            "Cài đặt kiosk trên thiết bị phần cứng (nếu có)",
            "Smoke test: đăng nhập, tạo phòng, tạo booking thử",
        ]),
        ("Giai Đoạn 2 — Khai Báo Dữ Liệu (Ngày 6–10)", [
            "Nhập thông tin khách sạn: tên, địa chỉ, MST, tài khoản ngân hàng",
            "Tạo loại phòng và upload ảnh từng loại (tối đa 10 ảnh/loại)",
            "Tạo phòng vật lý: số phòng, tầng, gán loại phòng",
            "Cấu hình gói giá: giá ngày thường, cuối tuần, discount",
            "Cấu hình thuế VAT và phí dịch vụ",
            "Tạo danh mục dịch vụ bán thêm (spa, breakfast, laundry...)",
            "Nhập nội dung trang Booking Portal: Giới thiệu, Tiện nghi, Liên hệ (VI + EN)",
            "Tạo tài khoản nhân viên và phân quyền",
            "Cấu hình email template xác nhận booking",
        ]),
        ("Giai Đoạn 3 — Training (Ngày 11–14)", [
            "Session 1 (2–3h): Admin — Cấu hình hệ thống, quản lý nhân viên",
            "Session 2 (2–3h): Lễ tân — Booking, check-in/out, folio, thanh toán",
            "Session 3 (1h): Housekeeping — Dashboard, Kanban board",
            "Session 4 (1h): Kế toán — Báo cáo, xuất hóa đơn PDF",
            "Hướng dẫn vận hành kiosk cho nhân viên trực quầy",
            "Cung cấp tài liệu hướng dẫn sử dụng",
        ]),
        ("Giai Đoạn 4 — UAT & Go-live (Ngày 14–21)", [
            "Chạy thử toàn bộ luồng: đặt phòng → kiosk → folio → báo cáo",
            "Nhân viên thực hành với dữ liệu thật (booking thực)",
            "Giải quyết feedback và điều chỉnh cấu hình",
            "Chuyển đổi từ hệ thống cũ (migration dữ liệu nếu cần)",
            "Go-live chính thức",
            "Theo dõi sát sao 1 tuần đầu — hotline ưu tiên",
        ]),
    ]

    for phase_title, items in phases:
        heading(doc, phase_title, 2)
        for item in items:
            bullet(doc, item)

    heading(doc, "3. Yêu Cầu Phía Khách Hàng", 1)
    simple_table(doc,
        ["Tùy chọn hosting", "Yêu cầu", "Chi phí ước tính"],
        [
            ["Cloud VPS (khuyến nghị)", "2 vCPU, 4GB RAM, 50GB SSD", "200.000–500.000 đ/tháng"],
            ["Cloud managed (Railway, Render)", "Đăng ký tài khoản", "300.000–700.000 đ/tháng"],
            ["Server tại chỗ", "PC/Server, UPS, internet ổn định", "Chi phí phần cứng một lần"],
        ],
        col_widths=[5.5, 6, 5]
    )

    heading(doc, "Phần Cứng Kiosk (Khách Hàng Tự Mua, HMS Tư Vấn)", 2)
    simple_table(doc,
        ["Thiết bị", "Yêu cầu tối thiểu", "Khuyến nghị"],
        [
            ["Android tablet", "Android 10+, màn hình 10\"", "Samsung Tab A8 hoặc tương đương"],
            ["Windows kiosk", "Windows 10, 4GB RAM, màn hình cảm ứng", "Mini PC + màn hình cảm ứng 15.6\""],
            ["Giá đỡ kiosk", "Phù hợp với thiết bị", "Giá đỡ inox chống trộm"],
        ],
        col_widths=[4, 6, 6.5]
    )

    heading(doc, "4. Hỗ Trợ Sau Triển Khai", 1)
    simple_table(doc,
        ["Kênh hỗ trợ", "Thời gian", "Dành cho"],
        [
            ["Zalo Group (dự án)", "Liên tục", "Câu hỏi nhanh, cập nhật hệ thống"],
            ["Email support", "Giờ hành chính", "Yêu cầu có tài liệu, thay đổi cấu hình"],
            ["Hotline (Professional+)", "Theo SLA gói", "Sự cố khẩn cấp"],
            ["Zoom / Google Meet", "Theo lịch hẹn", "Training bổ sung, review hàng tháng"],
        ],
        col_widths=[5, 4, 7.5]
    )
    body(doc, "Project Manager được chỉ định riêng — liên hệ trực tiếp trong suốt quá trình triển khai và 3 tháng đầu vận hành.")

    add_divider(doc)
    contact_placeholder(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════════════

DOCS = [
    ("01-tong-quan-san-pham.docx", doc1_overview),
    ("02-danh-sach-tinh-nang.docx", doc2_features),
    ("03-kich-ban-demo.docx",       doc3_demo),
    ("04-bang-gia.docx",            doc4_pricing),
    ("05-roi-business-case.docx",   doc5_roi),
    ("06-faq-ban-hang.docx",        doc6_faq),
    ("07-ke-hoach-trien-khai.docx", doc7_deployment),
]

if __name__ == "__main__":
    for filename, builder in DOCS:
        path = os.path.join(DOCS_DIR, filename)
        print(f"Generating {filename}...", end=" ", flush=True)
        doc = builder()
        doc.save(path)
        size = os.path.getsize(path)
        print(f"✓  {size // 1024} KB")
    print("\nAll documents generated.")
